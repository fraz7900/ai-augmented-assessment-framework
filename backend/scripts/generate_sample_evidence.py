"""Generate the binary sample-evidence documents used for MANUAL testing.

Not part of the application (see scripts/README.md).

Why generated rather than committed
-----------------------------------
`data/sample_evidence/` currently holds only a `.md` and a `.txt`, so
several ingestion paths had nothing a human could actually try them
against: PDF page-number citations (ADR-0042), PDF running-header
normalisation and sentence-boundary chunking (ADR-0055), OCR of a
scanned document (ADR-0055), and XLSX/CSV row and sheet provenance
(ADR-0041, ADR-0052). Notably, OCR -- the feature this branch is named
for -- could not be exercised by hand at all, because no image-only PDF
existed anywhere in the repo.

The files are generated instead of committed because `backend/conftest.py`
already states this repo's position: "Fixtures generate binary test
documents (DOCX, PDF) at test time rather than committing binary
fixtures to the repo, per notebooks/README.md's principle that anything
decision-relevant should be reproducible, not an opaque checked-in
artifact." A checked-in PDF is exactly such an artifact -- you cannot
review it in a diff. Output goes to `data/sample_evidence/generated/`,
which is gitignored.

Every document is synthetic. "Northwind Grid Utility" is not a real
organization and no real organizational, personnel, or system data
appears anywhere here, matching the existing samples and the rule in
`data/sample_evidence/README.md`.

Usage:
    python scripts/generate_sample_evidence.py
"""

from __future__ import annotations

import io
from pathlib import Path

import openpyxl
from docx import Document as DocxDocument
from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont

REPO_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = REPO_ROOT / "data" / "sample_evidence" / "generated"

SYNTHETIC_NOTICE = (
    "SYNTHETIC DOCUMENT. Fabricated for development and demo purposes only. "
    "Northwind Grid Utility is not a real organization."
)

# Prose deliberately written as whole sentences of varied length: chunk
# edges snap to sentence boundaries (ADR-0055), and text with no
# terminators would silently exercise only the word-boundary fallback.
SECTIONS: list[tuple[str, str]] = [
    (
        "Access Review",
        "Access rights are reviewed on a quarterly basis by system owners. Any account inactive "
        "for more than 30 days is disabled automatically. Reinstating a disabled account requires "
        "approval from the account owner's manager and the security team. Reviews are recorded in "
        "the access governance register and retained for three years.",
    ),
    (
        "Privileged Access",
        "Privileged accounts are issued only where a documented operational need exists. Each "
        "privileged session is logged and the log is forwarded to the central collector within "
        "five minutes. Shared administrative credentials are prohibited on OT systems. Vendor "
        "privileged access is time-bound and revoked at the end of each engagement.",
    ),
    (
        "Incident Response",
        "Suspected incidents are escalated to the on-call responder within one hour of detection. "
        "The responder classifies the incident and, where the classification is high, notifies the "
        "operations director directly. Post-incident reviews are held within ten business days. "
        "Lessons learned are tracked to closure by the security team.",
    ),
    (
        "Continuity and Recovery",
        "Recovery objectives are defined for every system supporting generation, transmission, or "
        "distribution. Backups are verified monthly by restoring a sample to an isolated "
        "environment. The continuity plan is exercised at least annually. Exercise findings are "
        "assigned owners and due dates before the exercise is closed.",
    ),
]


def _write_text_pdf(path: Path) -> None:
    """A multi-page PDF WITH a real text layer and a running footer.

    The footer is the point: it is what makes this document exercise
    ADR-0055's running-header normalisation rather than just chunking.
    """
    pdf = FPDF()
    # Without this, set_y(-25) trips FPDF's auto page break and the
    # footer lands on a page of its own, which is a different document
    # shape than the one worth testing.
    pdf.set_auto_page_break(auto=False)
    pdf.set_font("Helvetica", size=11)
    for index, (heading, body) in enumerate(SECTIONS, start=1):
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 13)
        # fpdf2's multi_cell leaves the cursor at the RIGHT margin, so x
        # is reset before each one; otherwise the next multi_cell(w=0)
        # has no width left and fpdf raises "Not enough horizontal space
        # to render a single character".
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 8, heading)
        pdf.set_font("Helvetica", size=11)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, body)
        pdf.set_y(-25)
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", size=9)
        # Page-numbered, so its signature differs per page -- the case
        # digit-insensitive detection exists for.
        pdf.cell(0, 6, f"Northwind Grid Utility - Security Policy - Page {index}")
    path.write_bytes(bytes(pdf.output()))


def _write_scanned_pdf(path: Path) -> None:
    """A PDF whose text exists ONLY as pixels -- no text layer at all.

    This is the one that makes OCR testable by hand. Rendered at a size
    and weight that a recogniser can actually read; deliberately not
    pathological, since the goal is to demonstrate the happy path.
    """
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 30)
        bold = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 36)
    except OSError:  # pragma: no cover - depends on the host's fonts
        font = ImageFont.load_default(size=30)
        bold = ImageFont.load_default(size=36)

    heading, body = SECTIONS[0]
    lines = [line.strip() for line in body.split(". ") if line.strip()]

    image = Image.new("RGB", (1400, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 50), f"{heading} (scanned copy)", fill="black", font=bold)
    for index, line in enumerate(lines):
        text = line if line.endswith(".") else f"{line}."
        draw.text((60, 150 + index * 70), text, fill="black", font=font)
    draw.text((60, 820), SYNTHETIC_NOTICE[:60], fill="black", font=font)

    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)

    pdf = FPDF()
    pdf.add_page()
    pdf.image(buffer, x=10, y=10, w=190)
    path.write_bytes(bytes(pdf.output()))


def _write_docx(path: Path) -> None:
    """Real heading styles, so this takes the structure-aware chunking
    path rather than the fixed-window one."""
    document = DocxDocument()
    document.add_heading("Northwind Grid Utility Security Policy", level=1)
    document.add_paragraph(SYNTHETIC_NOTICE)
    for heading, body in SECTIONS:
        document.add_heading(heading, level=1)
        document.add_paragraph(body)
    document.save(path)


def _asset_rows() -> list[list[str]]:
    return [
        ["Asset Name", "Owner", "Criticality", "Last Review"],
        ["Firewall-01", "NetOps", "High", "2026-04-02"],
        ["Switch-12", "NetOps", "Medium", "2026-04-02"],
        ["Historian-03", "OT Engineering", "High", "2026-03-18"],
        ["RTU-221", "Field Operations", "High", "2026-02-27"],
        ["Jumphost-02", "Security", "Medium", "2026-05-11"],
    ]


def _write_xlsx(path: Path) -> None:
    """Two sheets, so sheet_name provenance (ADR-0052) is visible rather
    than trivially constant."""
    workbook = openpyxl.Workbook()
    assets = workbook.active
    assets.title = "Assets"
    for row in _asset_rows():
        assets.append(row)

    vendors = workbook.create_sheet("Vendors")
    for row in [
        ["Vendor", "Service", "Access Level", "Contract End"],
        ["Acme Controls", "RTU maintenance", "Time-bound privileged", "2026-12-31"],
        ["Borealis Systems", "Historian support", "Read-only", "2027-06-30"],
    ]:
        vendors.append(row)

    workbook.save(path)


def _write_csv(path: Path) -> None:
    rows = _asset_rows()
    path.write_text("\n".join(",".join(cell for cell in row) for row in rows) + "\n")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [
        ("synthetic_security_policy.pdf", _write_text_pdf, "PDF, text layer + running footer"),
        ("synthetic_scanned_policy.pdf", _write_scanned_pdf, "image-only PDF (no text layer)"),
        ("synthetic_security_policy.docx", _write_docx, "DOCX with real heading styles"),
        ("synthetic_asset_inventory.xlsx", _write_xlsx, "XLSX, two sheets"),
        ("synthetic_asset_inventory.csv", _write_csv, "CSV"),
    ]
    for filename, writer, description in outputs:
        path = OUTPUT_DIR / filename
        writer(path)
        print(f"  {filename:38} {path.stat().st_size:>8,} bytes  {description}")
    print(f"\nwrote {len(outputs)} sample documents to {OUTPUT_DIR.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
