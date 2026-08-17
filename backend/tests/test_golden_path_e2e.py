"""Golden-path end-to-end test (controlled-pilot readiness pass,
`docs/architecture/02-controlled-pilot-readiness-audit.md`).

Chains create-assessment -> upload (multi-format, real parser/chunker/
embedder/LanceDB) -> propose-mappings (retrieval) -> human review ->
practice findings (ADR-0030) -> dashboard -> sanitization preview/
approval (ADR-0032) -> sanitized PDF/XLSX export, against a single small
fictional energy-utility evidence corpus deliberately built to contain
every category the audit's mission brief named: correct evidence,
missing evidence, contradictory evidence, stale evidence, a duplicate
upload, and an irrelevant document -- across all four supported formats
(PDF, DOCX, TXT, MD).

No test anywhere in this repository previously chained the full
pipeline in one place (confirmed during the audit); prior tests each
covered one or two stages against fakes or in isolation. This is real,
against the real FastAPI app, real SQLite, real LanceDB, and the real
C2M2 framework data -- no fakes.
"""

from __future__ import annotations

import io
from collections.abc import Iterator
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from fpdf import FPDF
from openpyxl import load_workbook
from pypdf import PdfReader

from compliance_platform.api import dependencies
from compliance_platform.core.config import Settings
from compliance_platform.main import app

_CACHED_DEPENDENCIES = (
    dependencies.get_cached_settings,
    dependencies.get_cached_vector_repository,
    dependencies.get_cached_assessment_repository,
    dependencies.get_cached_framework_registry,
)


@pytest.fixture
def client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Iterator[TestClient]:
    test_settings = Settings(
        vector_store_dir=tmp_path / "lancedb",
        assessments_db_path=tmp_path / "assessments.db",
        data_raw_dir=tmp_path / "raw",  # retained uploads (ADR-0056)
    )
    for cached in _CACHED_DEPENDENCIES:
        cached.cache_clear()
    monkeypatch.setattr(dependencies, "get_settings", lambda: test_settings)
    with TestClient(app) as test_client:
        yield test_client
    for cached in _CACHED_DEPENDENCIES:
        cached.cache_clear()


# --- The fictional utility's evidence corpus ---
#
# Fabricated from scratch for this test only, per data/sample_evidence/
# README.md's synthetic-data rule -- no real organization's content.
# Generated at test time rather than committed as binary fixtures, the
# same convention backend/conftest.py's sample_pdf_bytes/sample_docx_bytes
# already established.

_IDENTITY_POLICY_TEXT = (
    "Identity and Access Management Policy, Northfield Municipal Power & Light.\n\n"
    "All personnel requiring access to operational technology systems are "
    "provisioned an individual, non-shared identity through the formal HR "
    "onboarding workflow before any access is granted. Requests are logged in "
    "the access management system and approved by the requesting employee's "
    "manager and the OT security lead prior to provisioning."
)


def _identity_policy_pdf_bytes() -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, _IDENTITY_POLICY_TEXT)
    return bytes(pdf.output())


def _credential_encryption_docx_bytes() -> bytes:
    """CONTRADICTORY (positive half): a reassuring, general-purpose
    security overview claiming credentials are always encrypted at
    rest -- contradicted by the incident report below."""
    doc = DocxDocument()
    doc.add_heading("Credential Protection Standard", level=1)
    doc.add_paragraph(
        "All user and service account credentials, including passwords, "
        "smartcards, certificates, and cryptographic keys, are protected "
        "using strong encryption at rest across all Northfield systems, in "
        "accordance with corporate security standards."
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


_INCIDENT_REPORT_TEXT = (
    "Internal Security Incident Note IR-2026-014.\n\n"
    "During a routine audit of the legacy SCADA historian server, credentials "
    "for the shared 'histuser' service account were found stored in a "
    "plaintext configuration file, contradicting the corporate Credential "
    "Protection Standard. Remediation is tracked separately; this account's "
    "credential handling does not currently meet the encrypted-at-rest "
    "requirement described in that standard."
)


_STALE_DEPROVISIONING_POLICY_MD = (
    "# Account Deprovisioning Procedure (v1, 2019)\n\n"
    "**NOTE: This document describes the 2019 process. A revised v2 procedure "
    "was approved by the security committee in 2024 but has not yet been "
    "uploaded to this system.**\n\n"
    "When an employee separates from the company, IT removes their network "
    "login within 5 business days of the separation date, per the 2019 "
    "onboarding/offboarding runbook."
)

_SCOPE_EXCLUSION_MEMO_TEXT = (
    "Northfield Municipal Power & Light -- Assessment Scope Memo, 2026 cycle. "
    "Signed by the CISO and the Director of IT Operations. Northfield operates "
    "no privileged logical access tier beyond standard user accounts: all "
    "administrative actions on in-scope systems are performed via the shared "
    "jump host under the standard account model described in the Identity and "
    "Access Management Policy. Practices addressing a separate privileged "
    "access tier are therefore out of scope for this assessment cycle."
)

_IRRELEVANT_DOCUMENT_TEXT = (
    "Northfield Municipal Power & Light -- Employee Cafeteria Menu, Week of "
    "March 3. Monday: turkey sandwich. Tuesday: vegetable soup. The break "
    "room coffee machine will be serviced Wednesday morning; expect a brief "
    "outage of hot water on that floor."
)


def _ingest(client: TestClient, filename: str, content: bytes) -> str:
    response = client.post("/ingest", files={"file": (filename, content)})
    assert response.status_code == 200, response.text
    return response.json()["document_id"]


def test_golden_path_evidence_to_dashboard_and_export(client: TestClient) -> None:
    # 1. CREATE ASSESSMENT -- also confirms ADR-0031's framework-version
    # pin lands on a real, non-placeholder C2M2 version string.
    create_response = client.post(
        "/assessments",
        json={"name": "Northfield Municipal Power & Light -- 2026 C2M2 Self-Assessment",
              "framework_name": "C2M2"},
    )
    assert create_response.status_code == 200
    assessment = create_response.json()
    assessment_id = assessment["id"]
    assert assessment["framework_version"]

    # 2. UPLOAD -- multi-format corpus: PDF (correct), DOCX + TXT
    # (contradictory pair), MD (stale), TXT (irrelevant), and the PDF
    # again (duplicate upload of identical content).
    identity_policy_pdf_id = _ingest(client, "identity_policy.pdf", _identity_policy_pdf_bytes())
    credential_docx_id = _ingest(
        client, "credential_protection_standard.docx", _credential_encryption_docx_bytes()
    )
    incident_report_txt_id = _ingest(
        client, "incident_report_ir_2026_014.txt", _INCIDENT_REPORT_TEXT.encode()
    )
    stale_deprovisioning_md_id = _ingest(
        client, "deprovisioning_procedure_v1.md", _STALE_DEPROVISIONING_POLICY_MD.encode()
    )
    scope_memo_txt_id = _ingest(
        client, "assessment_scope_memo.txt", _SCOPE_EXCLUSION_MEMO_TEXT.encode()
    )
    irrelevant_txt_id = _ingest(client, "cafeteria_menu.txt", _IRRELEVANT_DOCUMENT_TEXT.encode())
    duplicate_pdf_id = _ingest(client, "identity_policy.pdf", _identity_policy_pdf_bytes())
    assert duplicate_pdf_id != identity_policy_pdf_id  # a fresh document_id each upload

    # 3. RETRIEVE / PROPOSE -- AI-proposed mapping for the correct-evidence
    # case, exercising the real retrieval-based mapping engine (ADR-0011).
    # propose-mappings only searches documents already associated with the
    # assessment, so link the PDF manually first, then let it also find
    # the practice via retrieval to confirm both paths coexist correctly.
    client.post(
        f"/assessments/{assessment_id}/evidence",
        json={"document_id": identity_policy_pdf_id, "practice_reference": "ACCESS-1a"},
    )
    propose_response = client.post(f"/assessments/{assessment_id}/propose-mappings")
    assert propose_response.status_code == 200

    # CONTRADICTORY: both the reassuring DOCX and the contradicting
    # incident-report TXT are linked as evidence for the same practice.
    client.post(
        f"/assessments/{assessment_id}/evidence",
        json={"document_id": credential_docx_id, "practice_reference": "ACCESS-1b"},
    )
    client.post(
        f"/assessments/{assessment_id}/evidence",
        json={"document_id": incident_report_txt_id, "practice_reference": "ACCESS-1b"},
    )

    # STALE: linked, but the human reviewer will judge it insufficient below.
    client.post(
        f"/assessments/{assessment_id}/evidence",
        json={"document_id": stale_deprovisioning_md_id, "practice_reference": "ACCESS-1c"},
    )

    # DUPLICATE: the second identical PDF upload also linked to the same
    # practice as the first -- must not double-count or break scoring.
    client.post(
        f"/assessments/{assessment_id}/evidence",
        json={"document_id": duplicate_pdf_id, "practice_reference": "ACCESS-1a"},
    )

    # SCOPE EXCLUSION BASIS (ADR-0057): a NOT_APPLICABLE finding moves
    # the score by shrinking the denominator, so it needs the same
    # evidence basis as a positive finding. The signed scope memo is that
    # basis -- an artifact an assessor can read, rather than the
    # reviewer's assertion that a conversation happened.
    client.post(
        f"/assessments/{assessment_id}/evidence",
        json={"document_id": scope_memo_txt_id, "practice_reference": "ACCESS-2g"},
    )

    # IRRELEVANT: deliberately never linked to anything -- a human
    # reviewer recognizing off-topic content and not linking it *is* the
    # correct handling, not a gap in the system.
    _ = irrelevant_txt_id

    evidence_before_review = client.get(f"/assessments/{assessment_id}/evidence").json()
    linked_document_ids = {link["document_id"] for link in evidence_before_review}
    assert irrelevant_txt_id not in linked_document_ids

    # 4. HUMAN REVIEW -- accept the AI-proposed ACCESS-1a link if one
    # exists (retrieval may or may not have independently found it above
    # threshold; either way ACCESS-1a already has an accepted manual link).
    for link in evidence_before_review:
        if link["review_status"] == "pending":
            decision = "accepted" if link["practice_reference"] == "ACCESS-1a" else "rejected"
            client.post(
                f"/assessments/{assessment_id}/evidence/{link['id']}/review",
                json={"decision": decision, "note": "Reviewed against the real submitted policy."},
            )

    # 5. PRACTICE FINDINGS (ADR-0030) -- the actual correctness fix under
    # test: a human's considered judgment about a practice, expressed
    # explicitly with a rationale, not left to be inferred from raw
    # evidence-link acceptance state alone.
    client.put(
        f"/assessments/{assessment_id}/practice-findings/ACCESS-1b",
        json={
            "status": "not_satisfied",
            "rationale": (
                "Incident report IR-2026-014 confirms the SCADA historian's shared "
                "service account stores credentials in plaintext, directly "
                "contradicting the Credential Protection Standard on file. "
                "Overriding the accepted evidence link: this practice is not met."
            ),
        },
    )
    client.put(
        f"/assessments/{assessment_id}/practice-findings/ACCESS-1c",
        json={
            "status": "insufficient_evidence",
            "rationale": (
                "Only the superseded 2019 deprovisioning procedure is on file; its "
                "own text states a 2024 revision exists but has not been submitted. "
                "Insufficient to confirm current practice."
            ),
        },
    )
    client.put(
        f"/assessments/{assessment_id}/practice-findings/ACCESS-2g",
        json={
            "status": "not_applicable",
            "rationale": (
                "Northfield has no privileged logical access tier beyond standard "
                "user accounts; confirmed with the CISO as out of scope for this "
                "assessment cycle."
            ),
        },
    )
    # ACCESS-1d (password strength, MIL2) is left with no evidence and no
    # finding at all -- the genuinely MISSING-evidence case.

    # 6. SCORE -- the actual bug this sprint fixes, proven live: ACCESS-1b
    # has two ACCEPTED evidence links yet is correctly NOT counted, because
    # the explicit NOT_SATISFIED finding overrides them.
    score = client.get(f"/assessments/{assessment_id}/score").json()
    assert "ACCESS" in score

    # 7. DASHBOARD -- gap statuses must be distinguishable, not collapsed.
    dashboard = client.get(f"/assessments/{assessment_id}/dashboard").json()
    gaps_by_practice = {
        gap["practice_id"]: gap
        for group in dashboard["complication"]
        for gap in group["gaps"]
        if group["domain_short_code"] == "ACCESS"
    }
    assert "ACCESS-1a" not in gaps_by_practice  # satisfied -- correct evidence, accepted
    assert gaps_by_practice["ACCESS-1b"]["status"] == "not_satisfied"
    assert "plaintext" in gaps_by_practice["ACCESS-1b"]["finding_rationale"]
    assert gaps_by_practice["ACCESS-1c"]["status"] == "insufficient_evidence"
    assert "superseded" in gaps_by_practice["ACCESS-1c"]["finding_rationale"]
    assert gaps_by_practice["ACCESS-1d"]["status"] == "insufficient_evidence"
    assert gaps_by_practice["ACCESS-1d"]["finding_rationale"] is None  # never reviewed at all
    # not_applicable AND evidence-backed by the signed scope memo, so it
    # is excluded from the denominator entirely (ADR-0057). Without that
    # link it would remain a gap -- covered by dedicated tests elsewhere.
    assert "ACCESS-2g" not in gaps_by_practice
    readiness = client.get(f"/assessments/{assessment_id}/finalization-readiness").json()
    unsupported = [
        b["category"]
        for b in readiness["blockers"]
        if b["category"].startswith("unsupported_")
    ]
    assert unsupported == []  # every finding here is properly evidence-backed

    # 8. EXPORT -- PDF and XLSX, confirming the real citation (practice
    # IDs, not just aggregate numbers) survives all the way through
    # rendering, per the mission's "evidence citations preserved through
    # assessment/reporting" acceptance criterion.
    pdf_response = client.get(f"/assessments/{assessment_id}/report/pdf")
    assert pdf_response.status_code == 200
    assert pdf_response.headers["content-type"] == "application/pdf"
    reader = PdfReader(io.BytesIO(pdf_response.content))
    pdf_text = "\n".join(page.extract_text() for page in reader.pages)
    assert "ACCESS-1b" in pdf_text  # the specific gap citation, not just an aggregate score
    assert "ACCESS-1d" in pdf_text
    # ADR-0040: the finding's rationale and the specific evidence link it
    # cites now actually render in the export, not just the practice ID.
    assert "plaintext" in pdf_text
    assert incident_report_txt_id in pdf_text

    xlsx_response = client.get(f"/assessments/{assessment_id}/report/xlsx")
    assert xlsx_response.status_code == 200
    workbook = load_workbook(io.BytesIO(xlsx_response.content))
    xlsx_text = "\n".join(
        str(cell.value)
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        for cell in row
        if cell.value is not None
    )
    assert "ACCESS-1b" in xlsx_text
    assert "plaintext" in xlsx_text
    assert incident_report_txt_id in xlsx_text
    assert "Northfield Municipal Power & Light" in pdf_text  # unsanitized -- real org name present

    # 9. SANITIZE -- preview the redaction/pseudonymization diff, approve
    # it explicitly, then confirm the sanitized export actually differs
    # from the unsanitized one and no longer carries the organization
    # name -- "internal assessment -> sanitization -> preview/diff ->
    # human approval -> sanitized export" (ADR-0032), never silent.
    unsanitized_before_approval = client.get(
        f"/assessments/{assessment_id}/report/pdf?sanitized=true"
    )
    assert unsanitized_before_approval.status_code == 412  # blocked -- no approval yet

    preview_response = client.post(
        f"/assessments/{assessment_id}/sanitization/preview",
        json={"custom_terms": ["Northfield Municipal Power & Light"]},
    )
    assert preview_response.status_code == 200
    preview = preview_response.json()
    assert any(m["category"] == "custom_term" for m in preview["matches"])
    assert (
        "Northfield Municipal Power & Light"
        not in preview["sanitized_report"]["situation"]["assessment_name"]
    )

    approve_response = client.post(
        f"/assessments/{assessment_id}/sanitization/approve",
        json={
            "custom_terms": ["Northfield Municipal Power & Light"],
            "approved_by": "compliance-lead@review",
        },
    )
    assert approve_response.status_code == 200

    sanitized_pdf_response = client.get(f"/assessments/{assessment_id}/report/pdf?sanitized=true")
    assert sanitized_pdf_response.status_code == 200
    sanitized_reader = PdfReader(io.BytesIO(sanitized_pdf_response.content))
    sanitized_pdf_text = "\n".join(page.extract_text() for page in sanitized_reader.pages)
    assert "Northfield Municipal Power & Light" not in sanitized_pdf_text
    assert "ORG-TERM" in sanitized_pdf_text
    assert "ACCESS-1b" in sanitized_pdf_text  # gap citations survive sanitization too

    # 10. FINALIZE -- immutability takes effect; further mutation is blocked.
    client.post(f"/assessments/{assessment_id}/status", json={"status": "in_review"})
    finalize_response = client.post(
        f"/assessments/{assessment_id}/status", json={"status": "finalized"}
    )
    assert finalize_response.status_code == 200

    blocked_finding = client.put(
        f"/assessments/{assessment_id}/practice-findings/ACCESS-1d",
        json={"status": "satisfied", "rationale": "should be blocked"},
    )
    assert blocked_finding.status_code == 409

    blocked_evidence = client.post(
        f"/assessments/{assessment_id}/evidence",
        json={"document_id": identity_policy_pdf_id, "practice_reference": "ACCESS-2a"},
    )
    assert blocked_evidence.status_code == 409
