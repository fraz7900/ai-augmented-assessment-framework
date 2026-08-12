from __future__ import annotations

import io
import zipfile

import openpyxl
import pytest
from docx import Document as DocxDocument
from fpdf import FPDF
from pypdf import PdfWriter

from compliance_platform.core.config import Settings
from compliance_platform.models.schemas import FileType, ParseStatus
from compliance_platform.services import document_parsers


def test_parse_plain_text_success() -> None:
    parsed = document_parsers.parse_document("notes.txt", b"Some real evidence content here.")
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert parsed.metadata.file_type == FileType.TXT
    assert "evidence" in parsed.raw_text


def test_parse_plain_text_empty() -> None:
    parsed = document_parsers.parse_document("empty.txt", b"   \n\n  ")
    assert parsed.parse_status == ParseStatus.EMPTY


def test_parse_plain_text_handles_invalid_utf8_gracefully(invalid_utf8_bytes: bytes) -> None:
    parsed = document_parsers.parse_document("weird.txt", invalid_utf8_bytes)
    # latin-1 fallback can decode any byte sequence, so this still succeeds,
    # but must carry a warning flagging the fallback for human review.
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert parsed.parse_warnings


def test_parse_markdown_dispatches_as_markdown() -> None:
    parsed = document_parsers.parse_document("policy.md", b"# Heading\nBody text.")
    assert parsed.metadata.file_type == FileType.MARKDOWN
    assert parsed.parse_status == ParseStatus.SUCCESS


def test_parse_docx_success(sample_docx_bytes: bytes) -> None:
    parsed = document_parsers.parse_document("policy.docx", sample_docx_bytes)
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert "# Sample Policy" in parsed.raw_text
    assert "# Second Section" in parsed.raw_text


def test_parse_pdf_success(sample_pdf_bytes: bytes) -> None:
    parsed = document_parsers.parse_document("policy.pdf", sample_pdf_bytes)
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert len(parsed.raw_text.strip()) > 0


def test_parse_pdf_detects_scanned_document(scanned_like_pdf_bytes: bytes) -> None:
    parsed = document_parsers.parse_document("scanned.pdf", scanned_like_pdf_bytes)
    assert parsed.parse_status == ParseStatus.UNSUPPORTED_SCANNED
    assert parsed.parse_warnings


def test_parse_unsupported_extension_raises() -> None:
    with pytest.raises(ValueError):
        document_parsers.parse_document("archive.zip", b"PK\x03\x04")


def test_content_hash_is_deterministic_but_document_id_is_not() -> None:
    content = b"identical content"
    p1 = document_parsers.parse_document("a.txt", content)
    p2 = document_parsers.parse_document("b.txt", content)
    assert p1.metadata.content_hash == p2.metadata.content_hash
    assert p1.metadata.document_id != p2.metadata.document_id


# --- Sprint 9: closing real, measured coverage gaps in the parsers'
# failure-mode branches — the document-parsing skill's whole point
# (distinguish "parsed but sparse" from "failed to parse") had no direct
# test for the "cannot even open the file" case in either format. ---


def test_parse_pdf_handles_malformed_content() -> None:
    parsed = document_parsers.parse_document("broken.pdf", b"not a real pdf at all")
    assert parsed.parse_status == ParseStatus.FAILED
    assert parsed.parse_warnings


def test_parse_pdf_handles_zero_pages() -> None:
    buffer = io.BytesIO()
    PdfWriter().write(buffer)  # a syntactically valid PDF with no pages at all
    parsed = document_parsers.parse_document("empty.pdf", buffer.getvalue())
    assert parsed.parse_status == ParseStatus.EMPTY


def test_parse_docx_handles_malformed_content() -> None:
    parsed = document_parsers.parse_document("broken.docx", b"not a real docx at all")
    assert parsed.parse_status == ParseStatus.FAILED
    assert parsed.parse_warnings


def test_parse_docx_with_only_whitespace_paragraphs_is_empty() -> None:
    doc = DocxDocument()
    doc.add_paragraph("   ")
    doc.add_paragraph("")
    buffer = io.BytesIO()
    doc.save(buffer)
    parsed = document_parsers.parse_document("blank.docx", buffer.getvalue())
    assert parsed.parse_status == ParseStatus.EMPTY


# --- Security hardening (controlled-pilot readiness audit §A.12):
# content-sniffing (a renamed file must not reach a parser expecting
# content it doesn't actually contain) and a decompression-bomb ceiling
# on DOCX (a ZIP archive whose central directory claims a huge
# uncompressed size, checked without decompressing a single byte). ---


def test_parse_pdf_extension_with_plain_text_content_is_caught_by_content_sniffing() -> None:
    parsed = document_parsers.parse_document("fake.pdf", b"Just plain text, not a PDF at all.")
    assert parsed.parse_status == ParseStatus.FAILED
    assert any("does not match its .pdf extension" in w for w in parsed.parse_warnings)


def test_parse_docx_extension_with_plain_text_content_is_caught_by_content_sniffing() -> None:
    parsed = document_parsers.parse_document("fake.docx", b"Just plain text, not a docx zip.")
    assert parsed.parse_status == ParseStatus.FAILED
    assert any("does not match its .docx extension" in w for w in parsed.parse_warnings)


def test_parse_txt_extension_with_binary_content_is_caught_by_content_sniffing() -> None:
    binary_content = b"some header\x00\x01\x02binary garbage that is not real text"
    parsed = document_parsers.parse_document("fake.txt", binary_content)
    assert parsed.parse_status == ParseStatus.FAILED
    assert any("does not match its .txt extension" in w for w in parsed.parse_warnings)


def test_parse_pdf_with_valid_signature_but_invalid_body_fails_in_parser_not_sniffing() -> None:
    # Passes the magic-byte check (proves sniffing and parser-level
    # failure are genuinely distinct code paths, not the same check
    # under two names) but is not a real, structurally valid PDF.
    parsed = document_parsers.parse_document(
        "fake2.pdf", b"%PDF-1.4\ngarbage body, not real PDF structure at all, no xref table"
    )
    assert parsed.parse_status == ParseStatus.FAILED
    assert not any("does not match its" in w for w in parsed.parse_warnings)
    assert any("Could not open PDF" in w for w in parsed.parse_warnings)


def test_parse_docx_rejects_a_zip_bomb_style_archive_before_decompressing() -> None:
    # A real ZIP with a genuinely high compression ratio (250MB of zero
    # bytes compresses to under 1MB) -- exactly the shape a decompression
    # bomb takes, not a synthetic/fabricated ZipInfo. The ceiling check
    # reads only the central directory's file_size metadata, so this
    # must reject in well under a second, never decompressing the payload.
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"\x00" * (250 * 1024 * 1024))

    parsed = document_parsers.parse_document("bomb.docx", buffer.getvalue())
    assert parsed.parse_status == ParseStatus.FAILED
    assert any("decompression-bomb" in w for w in parsed.parse_warnings)


# --- XLSX/CSV parsing (Sprint 18, ADR-0041) ---


def test_parse_xlsx_success_renders_sheet_heading_and_row_citations(
    sample_xlsx_bytes: bytes,
) -> None:
    parsed = document_parsers.parse_document("inventory.xlsx", sample_xlsx_bytes)
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert parsed.metadata.file_type == FileType.XLSX
    assert "# Assets" in parsed.raw_text  # sheet name as a heading, for structure-aware chunking
    assert "Row 2: Asset Name: Firewall-01 | Owner: NetOps | Criticality: High" in parsed.raw_text
    assert "Row 3: Asset Name: Switch-12 | Owner: NetOps | Criticality: Medium" in parsed.raw_text


def test_parse_xlsx_renders_multiple_sheets_as_separate_headings() -> None:
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Assets"
    ws1.append(["Asset Name"])
    ws1.append(["Firewall-01"])
    ws2 = wb.create_sheet("Vendors")
    ws2.append(["Vendor"])
    ws2.append(["Acme Corp"])
    buffer = io.BytesIO()
    wb.save(buffer)

    parsed = document_parsers.parse_document("inventory.xlsx", buffer.getvalue())
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert "# Assets" in parsed.raw_text
    assert "# Vendors" in parsed.raw_text
    assert parsed.raw_text.index("# Assets") < parsed.raw_text.index("# Vendors")


def test_parse_xlsx_skips_blank_rows() -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Asset Name"])
    ws.append(["Firewall-01"])
    ws.append([None])  # a fully blank row
    ws.append(["Switch-12"])
    buffer = io.BytesIO()
    wb.save(buffer)

    parsed = document_parsers.parse_document("inventory.xlsx", buffer.getvalue())
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert "Row 2:" in parsed.raw_text
    assert "Row 4:" in parsed.raw_text  # blank row 3 skipped, not rendered as an empty citation


def test_parse_xlsx_with_only_a_header_row_is_empty() -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Asset Name", "Owner"])
    buffer = io.BytesIO()
    wb.save(buffer)

    parsed = document_parsers.parse_document("inventory.xlsx", buffer.getvalue())
    assert parsed.parse_status == ParseStatus.EMPTY


def test_parse_xlsx_handles_malformed_content() -> None:
    parsed = document_parsers.parse_document(
        "broken.xlsx", b"PK\x03\x04not a real xlsx workbook at all"
    )
    assert parsed.parse_status == ParseStatus.FAILED
    assert parsed.parse_warnings


def test_parse_xlsx_rejects_a_zip_bomb_style_archive_before_decompressing() -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("xl/worksheets/sheet1.xml", b"\x00" * (250 * 1024 * 1024))

    parsed = document_parsers.parse_document("bomb.xlsx", buffer.getvalue())
    assert parsed.parse_status == ParseStatus.FAILED
    assert any("decompression-bomb" in w for w in parsed.parse_warnings)


def test_parse_xlsx_extension_with_plain_text_content_is_caught_by_content_sniffing() -> None:
    parsed = document_parsers.parse_document("fake.xlsx", b"Just plain text, not an xlsx zip.")
    assert parsed.parse_status == ParseStatus.FAILED
    assert any("does not match its .xlsx extension" in w for w in parsed.parse_warnings)


def test_parse_csv_success_renders_row_citations(sample_csv_bytes: bytes) -> None:
    parsed = document_parsers.parse_document("inventory.csv", sample_csv_bytes)
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert parsed.metadata.file_type == FileType.CSV
    assert "Row 2: Asset Name: Firewall-01 | Owner: NetOps | Criticality: High" in parsed.raw_text
    assert "Row 3: Asset Name: Switch-12 | Owner: NetOps | Criticality: Medium" in parsed.raw_text


def test_parse_csv_handles_ragged_rows_without_crashing() -> None:
    # A row with fewer columns than the header, and one with more --
    # real-world CSV exports are frequently ragged like this.
    content = b"Asset Name,Owner,Criticality\nFirewall-01,NetOps\nSwitch-12,NetOps,Medium,Extra\n"
    parsed = document_parsers.parse_document("inventory.csv", content)
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert "Row 2: Asset Name: Firewall-01 | Owner: NetOps" in parsed.raw_text
    assert "Row 3: Asset Name: Switch-12 | Owner: NetOps | Criticality: Medium" in parsed.raw_text


def test_parse_csv_with_only_a_header_row_is_empty() -> None:
    parsed = document_parsers.parse_document("inventory.csv", b"Asset Name,Owner\n")
    assert parsed.parse_status == ParseStatus.EMPTY


def test_parse_csv_handles_invalid_utf8_gracefully(invalid_utf8_bytes: bytes) -> None:
    content = b"Name,Note\n" + invalid_utf8_bytes.replace(b",", b";") + b"\n"
    parsed = document_parsers.parse_document("weird.csv", content)
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert parsed.parse_warnings  # latin-1 fallback warning, same convention as parse_plain_text


def test_parse_csv_extension_with_binary_content_is_caught_by_content_sniffing() -> None:
    binary_content = b"some header\x00\x01\x02binary garbage that is not real text"
    parsed = document_parsers.parse_document("fake.csv", binary_content)
    assert parsed.parse_status == ParseStatus.FAILED
    assert any("does not match its .csv extension" in w for w in parsed.parse_warnings)


# --- parser_version / page_number provenance (Sprint 18, ADR-0042) ---


def test_parser_version_reports_the_real_installed_pypdf_version(sample_pdf_bytes: bytes) -> None:
    """PDF provenance names pypdf AND this module's own normaliser: the
    stored text is pypdf's extraction *after* _normalise_page_texts ran,
    so pypdf's version alone would not identify what produced it.
    """
    import pypdf

    parsed = document_parsers.parse_document("policy.pdf", sample_pdf_bytes)
    assert parsed.metadata.parser_version == (
        f"pypdf=={pypdf.__version__}"
        f"+cp_pdf_normalizer=={document_parsers._PDF_NORMALIZER_VERSION}"
    )


def test_parser_version_reports_the_real_installed_python_docx_version(
    sample_docx_bytes: bytes,
) -> None:
    import importlib.metadata

    parsed = document_parsers.parse_document("policy.docx", sample_docx_bytes)
    expected_version = f"python-docx=={importlib.metadata.version('python-docx')}"
    assert parsed.metadata.parser_version == expected_version


def test_parser_version_reports_the_real_installed_openpyxl_version(
    sample_xlsx_bytes: bytes,
) -> None:
    parsed = document_parsers.parse_document("inventory.xlsx", sample_xlsx_bytes)
    assert parsed.metadata.parser_version == f"openpyxl=={openpyxl.__version__}"


def test_parser_version_for_stdlib_only_formats_reports_this_modules_own_version() -> None:
    for filename, content in [
        ("notes.txt", b"Some real evidence content here."),
        ("notes.md", b"# Heading\nBody text."),
        ("notes.csv", b"Name,Value\nA,1\n"),
    ]:
        parsed = document_parsers.parse_document(filename, content)
        assert parsed.metadata.parser_version.startswith("compliance_platform.document_parsers==")


def test_parse_pdf_returns_page_boundaries_that_slice_back_to_each_pages_real_text() -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, "Page one content. " * 5)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, "Page two content. " * 5)
    content = bytes(pdf.output())

    parsed = document_parsers.parse_document("multi.pdf", content)
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert parsed.page_boundaries is not None
    assert len(parsed.page_boundaries) == 2
    for start, end in parsed.page_boundaries:
        assert 0 <= start <= end <= len(parsed.raw_text)
    first_start, first_end = parsed.page_boundaries[0]
    assert "Page one content" in parsed.raw_text[first_start:first_end]
    second_start, second_end = parsed.page_boundaries[1]
    assert "Page two content" in parsed.raw_text[second_start:second_end]


def test_non_pdf_formats_have_no_page_boundaries(sample_docx_bytes: bytes) -> None:
    parsed = document_parsers.parse_document("policy.docx", sample_docx_bytes)
    assert parsed.page_boundaries is None


# --- row_boundaries (Sprint 18, ADR-0052) ---


def test_parse_xlsx_returns_row_boundaries_that_slice_back_to_each_rows_real_text() -> None:
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Assets"
    ws1.append(["Asset Name"])
    ws1.append(["Firewall-01"])
    ws1.append(["Switch-12"])
    ws2 = wb.create_sheet("Vendors")
    ws2.append(["Vendor"])
    ws2.append(["Acme Corp"])
    buffer = io.BytesIO()
    wb.save(buffer)

    parsed = document_parsers.parse_document("inventory.xlsx", buffer.getvalue())
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert parsed.page_boundaries is None
    assert parsed.row_boundaries is not None
    assert len(parsed.row_boundaries) == 3  # 2 Assets rows + 1 Vendors row
    for start, end, _row_number, _sheet_name in parsed.row_boundaries:
        assert 0 <= start <= end <= len(parsed.raw_text)

    by_sheet = {}
    for start, end, row_number, sheet_name in parsed.row_boundaries:
        by_sheet.setdefault(sheet_name, []).append((row_number, parsed.raw_text[start:end]))

    assert by_sheet["Assets"] == [
        (2, "Row 2: Asset Name: Firewall-01"),
        (3, "Row 3: Asset Name: Switch-12"),
    ]
    assert by_sheet["Vendors"] == [(2, "Row 2: Vendor: Acme Corp")]


def test_parse_csv_returns_row_boundaries_with_no_sheet_name(sample_csv_bytes: bytes) -> None:
    parsed = document_parsers.parse_document("inventory.csv", sample_csv_bytes)
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert parsed.page_boundaries is None
    assert parsed.row_boundaries is not None
    assert len(parsed.row_boundaries) == 2
    for start, end, row_number, sheet_name in parsed.row_boundaries:
        assert sheet_name is None  # CSV has no sheet concept
        assert f"Row {row_number}:" in parsed.raw_text[start:end]


def test_non_tabular_formats_have_no_row_boundaries(sample_docx_bytes: bytes) -> None:
    parsed = document_parsers.parse_document("policy.docx", sample_docx_bytes)
    assert parsed.row_boundaries is None


# --- PDF text normalisation: running headers/footers and blank runs ---
# (the defect ADR-0054 found while fixing chunk edges and deliberately
# left for a separate change, since it is extraction, not chunking)


# Body text must differ in WORDS between pages, not just in a page
# number: _running_line_signature deliberately erases digits, so
# "Section 1 content" and "Section 2 content" are one signature and are
# genuinely indistinguishable from a running footer. Using such text here
# would test the fallback guard rather than footer detection.
_PAGE_BODIES = [
    "Access control reviews are performed by the security team each quarter. ",
    "Incident response duties are assigned to named on-call engineering staff. ",
    "Vendor risk assessments are completed before any new contract is signed. ",
    "Backup restoration is exercised annually against documented recovery targets. ",
    "Change management approvals are recorded before production deployment. ",
]


def _footered_pdf_bytes(page_count: int, footer: str = "Incident Response Plan") -> bytes:
    """A multi-page PDF with real body text and a page-numbered running
    footer, i.e. the exact shape that produced ADR-0054's "'Plan 9' plus
    thirty newlines" chunks.
    """
    pdf = FPDF()
    # Without this, set_y(-25) trips FPDF's auto page break and every
    # footer lands on a page of its own -- which is a different document
    # shape than the one under test (a footer sharing a page with body
    # text), and quietly halves the footer's per-page hit rate.
    pdf.set_auto_page_break(auto=False)
    pdf.set_font("Helvetica", size=12)
    for page in range(1, page_count + 1):
        pdf.add_page()
        pdf.multi_cell(0, 10, _PAGE_BODIES[(page - 1) % len(_PAGE_BODIES)] * 3)
        pdf.set_y(-25)
        pdf.cell(0, 10, f"{footer} {page}")
    return bytes(pdf.output())


def test_running_page_footer_is_stripped_from_extracted_pdf_text() -> None:
    parsed = document_parsers.parse_document("policy.pdf", _footered_pdf_bytes(4))
    assert parsed.parse_status == ParseStatus.SUCCESS
    assert "Incident Response Plan" not in parsed.raw_text
    # The body text the footer sat next to is untouched.
    assert "Access control reviews" in parsed.raw_text
    assert "Vendor risk assessments" in parsed.raw_text


def test_page_boundaries_still_reconstruct_raw_text_after_normalisation() -> None:
    """The ADR-0042 invariant, re-asserted under normalisation: offsets
    describe the stored text exactly. Normalising after computing
    boundaries (rather than before) would silently break this.
    """
    parsed = document_parsers.parse_document("policy.pdf", _footered_pdf_bytes(4))
    assert parsed.page_boundaries is not None
    rebuilt = "\n\n".join(parsed.raw_text[start:end] for start, end in parsed.page_boundaries)
    assert rebuilt == parsed.raw_text


def test_blank_line_runs_are_collapsed_to_a_single_blank_line() -> None:
    # Closing lines differ per page: an identical closing line on every
    # page is a running footer by definition and would be stripped, which
    # is a different behaviour than the one under test here.
    pages = [f"{body}\n" + "\n" * 30 + f"Closing note about {body.split()[0].lower()}."
             for body in _PAGE_BODIES[:3]]
    normalised = document_parsers._normalise_page_texts(pages)
    for page in normalised:
        assert "\n\n\n" not in page
        assert "Closing note about" in page


def test_footer_differing_only_by_page_number_is_still_detected_as_running() -> None:
    pages = [f"{body}\n\nAnnual Security Policy {n}" for n, body in enumerate(_PAGE_BODIES, 1)]
    normalised = document_parsers._normalise_page_texts(pages)
    assert not any("Annual Security Policy" in page for page in normalised)
    assert all(body.strip() in page for body, page in zip(_PAGE_BODIES, normalised, strict=True))


def test_a_page_is_never_emptied_by_running_line_removal() -> None:
    """The content-loss fallback. Three identical short pages make every
    line look like a running header; keeping them beats deleting the
    document.
    """
    pages = ["Signature: ______\n\nDate: ______"] * 4
    normalised = document_parsers._normalise_page_texts(pages)
    assert all("Signature" in page for page in normalised)


def test_a_long_repeated_paragraph_is_never_treated_as_a_running_header() -> None:
    """The content-loss guard. Repeated boilerplate that is genuinely long
    is real evidence text -- deleting it would be a worse bug than the
    noise this normalisation removes.
    """
    long_line = (
        "This policy applies to all information systems owned or operated by the organization "
        "and to every third party granted access to those systems under contract."
    )
    assert len(long_line) > document_parsers._HEADER_FOOTER_MAX_CHARS
    pages = [f"{long_line}\n\nPage-specific body {n}." for n in range(1, 5)]
    normalised = document_parsers._normalise_page_texts(pages)
    assert all(long_line in page for page in normalised)


def test_a_line_repeated_in_the_middle_of_pages_is_kept() -> None:
    """Only the edge region is header/footer territory. A repeated line
    surrounded by other content is body text (a table row, a repeated
    clause) and must survive.
    """
    pages = [
        "\n".join(
            [f"Opening line {n}", "alpha", "beta", "REPEATED MIDDLE LINE", "gamma", "delta",
             f"Closing line {n}"]
        )
        for n in range(1, 5)
    ]
    normalised = document_parsers._normalise_page_texts(pages)
    assert all("REPEATED MIDDLE LINE" in page for page in normalised)


# --- OCR for scanned/image-only PDFs (ADR-0055) ---


def test_scanned_pdf_text_is_recovered_by_local_ocr(scanned_image_pdf_bytes: bytes) -> None:
    """The one test that runs the real recogniser end to end. Slow by
    nature (model load plus per-page inference), so the surrounding
    behaviours are covered by the stubbed tests below rather than by
    repeating a real OCR pass.
    """
    parsed = document_parsers.parse_document("scan.pdf", scanned_image_pdf_bytes)

    assert parsed.parse_status == ParseStatus.SUCCESS_OCR
    assert "privileged accounts" in parsed.raw_text
    assert "escalated" in parsed.raw_text
    # Provenance names what actually produced the text. pypdf did not.
    assert "rapidocr-onnxruntime==" in parsed.metadata.parser_version
    assert "pypdf==" not in parsed.metadata.parser_version
    assert any("recovered by local OCR" in w for w in parsed.parse_warnings)


def test_ocr_output_keeps_page_boundaries_consistent_with_stored_text(
    monkeypatch: pytest.MonkeyPatch, scanned_like_pdf_bytes: bytes
) -> None:
    from compliance_platform.services import ocr

    monkeypatch.setattr(
        ocr,
        "ocr_pdf_page_texts",
        lambda content, **kwargs: (["Recovered page text that is long enough to keep."], []),
    )
    parsed = document_parsers.parse_document("scan.pdf", scanned_like_pdf_bytes)

    assert parsed.parse_status == ParseStatus.SUCCESS_OCR
    assert parsed.page_boundaries is not None
    rebuilt = "\n\n".join(parsed.raw_text[start:end] for start, end in parsed.page_boundaries)
    assert rebuilt == parsed.raw_text


def test_ocr_disabled_reports_the_document_as_unsupported_scanned(
    scanned_image_pdf_bytes: bytes,
) -> None:
    settings = Settings(ocr_enabled=False)  # type: ignore[call-arg]
    parsed = document_parsers.parse_document(
        "scan.pdf", scanned_image_pdf_bytes, settings=settings
    )
    assert parsed.parse_status == ParseStatus.UNSUPPORTED_SCANNED
    assert any("OCR is disabled by configuration" in w for w in parsed.parse_warnings)


def test_missing_ocr_dependencies_fall_back_to_unsupported_scanned(
    monkeypatch: pytest.MonkeyPatch, scanned_like_pdf_bytes: bytes
) -> None:
    """An install without the optional OCR extras must degrade to the
    pre-OCR behaviour, not crash an upload.
    """
    from compliance_platform.services import ocr

    def _unavailable(content: bytes, **kwargs: object) -> tuple[list[str], list[str]]:
        raise ocr.OcrUnavailableError("rapidocr-onnxruntime is not installed")

    monkeypatch.setattr(ocr, "ocr_pdf_page_texts", _unavailable)
    parsed = document_parsers.parse_document("scan.pdf", scanned_like_pdf_bytes)
    assert parsed.parse_status == ParseStatus.UNSUPPORTED_SCANNED


def test_ocr_that_recovers_almost_nothing_is_still_unsupported_scanned(
    monkeypatch: pytest.MonkeyPatch, scanned_like_pdf_bytes: bytes
) -> None:
    """A near-empty OCR result must not be stored as though the document
    parsed -- that is precisely the "passed through as if successfully
    parsed" failure the document-parsing rule forbids.
    """
    from compliance_platform.services import ocr

    monkeypatch.setattr(ocr, "ocr_pdf_page_texts", lambda content, **kwargs: (["x"], []))
    parsed = document_parsers.parse_document("scan.pdf", scanned_like_pdf_bytes)
    assert parsed.parse_status == ParseStatus.UNSUPPORTED_SCANNED
    assert any("did not recover usable text" in w for w in parsed.parse_warnings)


def test_a_pdf_with_a_real_text_layer_never_invokes_ocr(
    monkeypatch: pytest.MonkeyPatch, sample_pdf_bytes: bytes
) -> None:
    """OCR costs seconds per page; the common case must not pay for it."""
    from compliance_platform.services import ocr

    def _fail(content: bytes, **kwargs: object) -> tuple[list[str], list[str]]:
        raise AssertionError("OCR must not run for a PDF that has a text layer")

    monkeypatch.setattr(ocr, "ocr_pdf_page_texts", _fail)
    parsed = document_parsers.parse_document("policy.pdf", sample_pdf_bytes)
    assert parsed.parse_status == ParseStatus.SUCCESS


def test_running_line_detection_requires_at_least_three_pages() -> None:
    """Two pages sharing a line is not evidence of a running footer; it
    could be a two-page document that legitimately repeats a heading.
    """
    pages = ["Body one.\n\nShared Footer 1", "Body two.\n\nShared Footer 2"]
    normalised = document_parsers._normalise_page_texts(pages)
    assert all("Shared Footer" in page for page in normalised)
