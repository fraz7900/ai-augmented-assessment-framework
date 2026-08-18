"""Document parsers: PDF, DOCX, TXT/Markdown, XLSX/CSV.

Implements the failure-mode handling required by the document-parsing
skill: a parser must distinguish "parsed successfully but short/sparse"
from "failed to parse, output is garbage" rather than always returning a
string and calling it done.
"""

from __future__ import annotations

import csv
import hashlib
import importlib.metadata
import io
import logging
import math
import re
import uuid
import zipfile
from collections import Counter

import openpyxl
from docx import Document as DocxDocument
from pypdf import PdfReader

from compliance_platform.core.config import Settings
from compliance_platform.models.schemas import (
    FileType,
    ParsedDocument,
    ParseStatus,
    SourceDocumentMetadata,
)

_logger = logging.getLogger(__name__)

# Reproducibility provenance (Sprint 18, ADR-0042): the real, installed
# version of whichever library actually parsed a document, read via
# importlib.metadata -- never hand-typed, so it can't drift out of sync
# with what's actually installed. TXT/MD/CSV have no third-party parsing
# library backing them (just this module's own decode/split logic), so
# they report this module's own hand-maintained version instead -- bump
# _PARSER_MODULE_VERSION when parse_plain_text/parse_csv/
# _render_tabular_rows's actual parsing behavior changes materially.
_PARSER_MODULE_VERSION = "1"

# (extracted text, status, warnings, page boundaries as (char_start, char_end)
# per page or None for formats with no page concept, row boundaries as
# (char_start, char_end, row_number, sheet_name) per rendered row or None
# for formats with no row concept) -- shared by every parse_*() function
# below. Exactly one of page/row boundaries is non-None for any given
# format (PDF has pages, XLSX/CSV have rows, everything else has neither).
ParseResult = tuple[
    str,
    ParseStatus,
    list[str],
    list[tuple[int, int]] | None,
    list[tuple[int, int, int, str | None]] | None,
]


# PDF text is normalised by this module after pypdf extracts it (see
# _normalise_page_texts), so "which pypdf version parsed this" is no
# longer the whole provenance story for a PDF chunk -- the normaliser's
# own version is part of what produced the stored text. Tracked
# separately from _PARSER_MODULE_VERSION so that changing PDF
# normalisation does not falsely signal a behaviour change in the
# TXT/MD/CSV parsers, which share nothing with it. Bump when
# _normalise_page_texts's actual output changes.
_PDF_NORMALIZER_VERSION = "1"


def _parser_version(file_type: FileType, *, status: ParseStatus | None = None) -> str:
    if file_type == FileType.PDF:
        if status == ParseStatus.SUCCESS_OCR:
            # pypdf never produced this text -- pdfium rendered the page
            # and the OCR recogniser read it. Naming pypdf here would be
            # wrong, not just incomplete.
            from compliance_platform.services import ocr

            return f"{ocr.engine_version()}+cp_pdf_normalizer=={_PDF_NORMALIZER_VERSION}"
        if status == ParseStatus.SUCCESS_PARTIAL_OCR:
            # Both really did produce text in this document, so both are
            # named. Recording only the recogniser would misattribute
            # every page that had a real text layer; recording only pypdf
            # would hide that some pages are approximate -- and
            # parser_version is what ADR-0042 makes citations reproducible
            # against.
            from compliance_platform.services import ocr

            return (
                f"pypdf=={importlib.metadata.version('pypdf')}"
                f"+{ocr.engine_version()}"
                f"+cp_pdf_normalizer=={_PDF_NORMALIZER_VERSION}"
            )
        return (
            f"pypdf=={importlib.metadata.version('pypdf')}"
            f"+cp_pdf_normalizer=={_PDF_NORMALIZER_VERSION}"
        )
    if file_type == FileType.DOCX:
        return f"python-docx=={importlib.metadata.version('python-docx')}"
    if file_type == FileType.XLSX:
        return f"openpyxl=={importlib.metadata.version('openpyxl')}"
    return f"compliance_platform.document_parsers=={_PARSER_MODULE_VERSION}"

# Below this many characters per page, a "successfully parsed" PDF is
# almost certainly a scanned/image-only document that happened to extract
# a few stray characters (e.g. a running header), not real text. Such a
# document is handed to OCR (ADR-0055) and only reported as
# UNSUPPORTED_SCANNED if that also fails to recover usable text. The same
# threshold is reused to judge the OCR output, so "usable" means the same
# thing on both paths.
_MIN_CHARS_PER_PAGE = 20

# Content-sniffing (controlled-pilot readiness audit §A.12, security
# hardening): file-type validation was extension-only before this —
# a file renamed to a different extension went straight to that
# extension's parser regardless of its real content. Checked against
# real magic bytes / a binary-content heuristic before any format-specific
# parser ever sees the bytes, not just trusted from the filename.
_PDF_SIGNATURE = b"%PDF-"
# DOCX and XLSX (both OOXML) are ZIP archives; PK\x03\x04 is a normal
# non-empty archive, PK\x05\x06 an empty one — both are genuine zip
# signatures.
_ZIP_SIGNATURES = (b"PK\x03\x04", b"PK\x05\x06")
_BINARY_SNIFF_BYTES = 8192


def _looks_like_binary_content(content: bytes) -> bool:
    # TXT/MD/CSV have no magic-byte signature to check; a NUL byte within
    # the first few KB is a standard, low-false-positive signal that this
    # is not real text content, however it got its extension. Real
    # invalid-UTF-8 text (mojibake, stray high bytes) does not typically
    # contain NUL bytes, so this does not fight the existing latin-1
    # fallback path in parse_plain_text/parse_csv.
    return b"\x00" in content[:_BINARY_SNIFF_BYTES]


def _content_matches_file_type(content: bytes, file_type: FileType) -> bool:
    if file_type == FileType.PDF:
        return content.startswith(_PDF_SIGNATURE)
    if file_type in (FileType.DOCX, FileType.XLSX):
        return content.startswith(_ZIP_SIGNATURES)
    if file_type in (FileType.TXT, FileType.MARKDOWN, FileType.CSV):
        return not _looks_like_binary_content(content)
    return True


# Decompression-bomb ceiling (same audit finding): DOCX and XLSX are
# both ZIP archives, and ZipInfo.file_size is real uncompressed-size
# metadata from the archive's own central directory -- readable without
# decompressing a single byte of entry content, so this check costs
# nothing close to what an actual decompression would. 200MB is far
# beyond any real policy document or spreadsheet's extracted size but
# well below what a crafted archive could claim to decompress to from a
# small uploaded file. CSV/TXT/MD are uncompressed plain text, already
# bounded directly by Settings.max_upload_bytes -- no separate ceiling
# needed for them.
_MAX_ZIP_UNCOMPRESSED_BYTES = 200 * 1024 * 1024


def _zip_bomb_ceiling_warning(content: bytes, format_label: str) -> tuple[int, str | None]:
    """Returns (total_uncompressed_bytes, warning). warning is None if
    within the ceiling (or the archive can't even be opened as a zip,
    which the caller's own DocxDocument()/load_workbook() call will
    separately and more specifically report as FAILED).
    """
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        total = sum(info.file_size for info in archive.infolist())
    if total > _MAX_ZIP_UNCOMPRESSED_BYTES:
        return total, (
            f"{format_label} would decompress to {total} bytes, exceeding the "
            f"{_MAX_ZIP_UNCOMPRESSED_BYTES}-byte safety ceiling. Rejected before extraction "
            "to guard against a decompression-bomb-style malformed file."
        )
    return total, None


def _content_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _new_document_id() -> str:
    return str(uuid.uuid4())


# PDF text normalisation (ADR-0054 disclosed this as a real, separate
# defect and deliberately left it: chunk edges were word-aligned, but the
# text between them still carried running headers/footers and long runs
# of blank lines, e.g. "Plan 9" followed by ~30 newlines). A chunk made
# mostly of a repeated page footer is word-aligned nonsense -- it embeds
# poorly, and when a reviewer opens it as cited evidence it says nothing
# about the control it was retrieved for.
#
# Only the first/last few non-blank lines of each page are candidates,
# because that is where running headers and footers live; a repeated line
# in the middle of a page is body text (a table row, a repeated clause)
# and is left alone.
_HEADER_FOOTER_EDGE_LINES = 3
# Fewer than 3 pages is not enough evidence that a line is *running*
# rather than coincidentally repeated.
_HEADER_FOOTER_MIN_PAGES = 3
# A line must appear in the edge region of at least this share of pages.
_HEADER_FOOTER_MIN_RATIO = 0.6
# Running headers/footers are short. This ceiling is the main guard
# against deleting real content: a long paragraph that happens to repeat
# across pages (boilerplate legal text, a repeated definition block) is
# never treated as a header, because losing it would be a genuine
# evidence loss rather than noise removal.
_HEADER_FOOTER_MAX_CHARS = 120

_BLANK_RUN_RE = re.compile(r"\n{3,}")
_DIGIT_RUN_RE = re.compile(r"\d+")


def _running_line_signature(line: str) -> str:
    """Collapse a line to a repeat-detection key: whitespace normalised,
    digit runs replaced, case folded. "Incident Response Plan 9" and
    "Incident Response Plan 10" share one signature, which is the whole
    point -- a page footer differs by its page number on every page and
    would otherwise never look repeated.
    """
    return _DIGIT_RUN_RE.sub("#", " ".join(line.split())).strip().lower()


def _edge_line_indices(lines: list[str]) -> set[int]:
    """Indices of the first and last few NON-BLANK lines. Computed over
    non-blank lines specifically so a page whose footer is preceded by
    blank lines still has that footer counted as an edge line -- which is
    exactly the case this normalisation exists to handle.
    """
    non_blank = [i for i, line in enumerate(lines) if line.strip()]
    return set(non_blank[:_HEADER_FOOTER_EDGE_LINES] + non_blank[-_HEADER_FOOTER_EDGE_LINES:])


def _detect_running_lines(page_texts: list[str]) -> set[str]:
    """Signatures of lines that recur in the edge region of most pages."""
    if len(page_texts) < _HEADER_FOOTER_MIN_PAGES:
        return set()

    counts: Counter[str] = Counter()
    for page_text in page_texts:
        lines = page_text.splitlines()
        edge = _edge_line_indices(lines)
        # A set per page: a header repeated twice on one page still only
        # counts as one page's worth of evidence that it is running.
        signatures = {
            _running_line_signature(lines[i])
            for i in edge
            if len(lines[i].strip()) <= _HEADER_FOOTER_MAX_CHARS
        }
        counts.update(sig for sig in signatures if sig)

    threshold = max(
        _HEADER_FOOTER_MIN_PAGES, math.ceil(len(page_texts) * _HEADER_FOOTER_MIN_RATIO)
    )
    return {sig for sig, seen_on in counts.items() if seen_on >= threshold}


def _normalise_page_text(page_text: str, running: set[str]) -> str:
    """Drop this page's running header/footer lines, strip trailing
    whitespace, and collapse blank-line runs to a single blank line.

    If dropping running lines would empty the page, nothing is dropped.
    On a short page every line is an "edge" line, so a document whose
    pages genuinely repeat each other (a form, a signature sheet, a
    two-line-per-page appendix) could otherwise have whole pages deleted
    as though they were headers. Losing a page of real evidence is a far
    worse failure than leaving a footer in one, so this fallback is
    unconditional rather than heuristic.
    """
    lines = page_text.splitlines()
    edge = _edge_line_indices(lines)
    kept = [
        line.rstrip()
        for i, line in enumerate(lines)
        if not (i in edge and line.strip() and _running_line_signature(line) in running)
    ]
    if page_text.strip() and not any(line.strip() for line in kept):
        kept = [line.rstrip() for line in lines]
    return _BLANK_RUN_RE.sub("\n\n", "\n".join(kept)).strip()


def _normalise_page_texts(page_texts: list[str]) -> list[str]:
    """Normalise every page, detecting running lines across the whole
    document first (a header is only identifiable by comparing pages).

    Returns normalised page texts rather than a joined string on purpose:
    parse_pdf builds page_boundaries from exactly the list it joins, so
    normalising before that step keeps offsets and text in sync by
    construction rather than by a second, drift-prone calculation.
    """
    running = _detect_running_lines(page_texts)
    return [_normalise_page_text(page_text, running) for page_text in page_texts]


def _page_boundaries_for(page_texts: list[str]) -> list[tuple[int, int]]:
    """Char offsets into "\\n\\n".join(page_texts) for each page.

    Shared by the text-layer and OCR paths so the two can never compute
    offsets differently -- the ADR-0042 invariant is that boundaries are
    derived from exactly the list that gets joined, and having that
    arithmetic in one place is what keeps it true.
    """
    boundaries: list[tuple[int, int]] = []
    cursor = 0
    for page_text in page_texts:
        start = cursor
        end = start + len(page_text)
        boundaries.append((start, end))
        cursor = end + 2  # the "\n\n" separator between pages
    return boundaries


def _ocr_pdf(content: bytes, page_count: int, settings: Settings) -> ParseResult | None:
    """Attempt OCR on a PDF with no usable text layer.

    Returns None when OCR is unavailable, disabled, or produced too
    little text to be worth storing -- in every one of those cases the
    caller falls back to UNSUPPORTED_SCANNED, i.e. to the exact behaviour
    that shipped before OCR existed.
    """
    from compliance_platform.services import ocr  # deferred: loads ONNX models

    try:
        page_texts, warnings = ocr.ocr_pdf_page_texts(
            content,
            dpi=settings.ocr_render_dpi,
            min_confidence=settings.ocr_min_confidence,
            max_pages=settings.ocr_max_pages,
        )
    except ocr.OcrUnavailableError as exc:
        _logger.warning("ocr unavailable: %s", exc)
        return None

    page_texts = _normalise_page_texts(page_texts)
    text = "\n\n".join(page_texts)
    if len(text) / page_count < _MIN_CHARS_PER_PAGE:
        return None

    warnings.append(
        "Text was recovered by local OCR because this PDF has no text layer. OCR output is "
        "approximate -- verify any passage against the source page before relying on it as "
        "evidence."
    )
    return text, ParseStatus.SUCCESS_OCR, warnings, _page_boundaries_for(page_texts), None


def _ocr_missing_pages(
    content: bytes,
    page_texts: list[str],
    textless: list[int],
    settings: Settings,
) -> tuple[list[str], list[int], list[str]] | None:
    """OCR only the pages pypdf found no text on, splicing the results
    into `page_texts`.

    Returns (merged_page_texts, recovered_page_indices, warnings), or
    None when OCR is unavailable or recovered nothing -- in which case
    the caller keeps the pypdf text exactly as it was. A document that
    parsed acceptably must never become worse for having tried OCR.
    """
    from compliance_platform.services import ocr  # deferred: loads ONNX models

    try:
        ocr_page_texts, warnings = ocr.ocr_pdf_page_texts(
            content,
            dpi=settings.ocr_render_dpi,
            min_confidence=settings.ocr_min_confidence,
            max_pages=settings.ocr_max_pages,
            only_pages=set(textless),
        )
    except ocr.OcrUnavailableError as exc:
        _logger.warning("ocr unavailable for partial-scan recovery: %s", exc)
        return None

    ocr_page_texts = _normalise_page_texts(ocr_page_texts)

    merged = list(page_texts)
    recovered: list[int] = []
    for index in textless:
        if index >= len(ocr_page_texts):
            continue
        candidate = ocr_page_texts[index]
        # Only replace when OCR actually did better than what was there.
        # Recognising a genuinely blank page yields noise or nothing, and
        # substituting that for an honestly-empty page would manufacture
        # evidence text out of a blank sheet.
        if len(candidate) >= _MIN_CHARS_PER_PAGE and len(candidate) > len(merged[index]):
            merged[index] = candidate
            recovered.append(index)

    if not recovered:
        return None
    return merged, recovered, warnings


def parse_pdf(content: bytes, settings: Settings | None = None) -> ParseResult:
    settings = settings if settings is not None else Settings()
    warnings: list[str] = []
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:  # pypdf raises varied exception types on malformed PDFs
        return "", ParseStatus.FAILED, [f"Could not open PDF: {exc}"], None, None

    page_count = len(reader.pages)
    if page_count == 0:
        return "", ParseStatus.EMPTY, ["PDF has zero pages."], None, None

    page_texts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            page_texts.append(page.extract_text() or "")
        except Exception as exc:
            warnings.append(f"Failed to extract text from page {i + 1}: {exc}")
            page_texts.append("")

    # Normalise BEFORE joining and before computing page_boundaries, so
    # offsets describe the text that is actually stored and cited, not
    # the pre-normalisation text that no longer exists anywhere.
    page_texts = _normalise_page_texts(page_texts)

    text = "\n\n".join(page_texts)

    # Page boundaries (Sprint 18, ADR-0042): char offsets into `text`
    # for each page, computed from the exact same page_texts/"\n\n".join
    # this function already builds, so they can never drift out of sync
    # with the actual joined text. Discarded before this sprint —
    # controlled-pilot readiness audit §A.3.
    page_boundaries = _page_boundaries_for(page_texts)

    avg_chars_per_page = len(text) / page_count

    if avg_chars_per_page < _MIN_CHARS_PER_PAGE:
        warnings.append(
            f"Average {avg_chars_per_page:.1f} extracted characters per page "
            f"across {page_count} page(s); this looks like a scanned or image-only PDF."
        )
        if settings.ocr_enabled:
            ocr_result = _ocr_pdf(content, page_count, settings)
            if ocr_result is not None:
                ocr_text, status, ocr_warnings, ocr_boundaries, _ = ocr_result
                return ocr_text, status, warnings + ocr_warnings, ocr_boundaries, None
            warnings.append(
                "OCR was attempted and did not recover usable text; treating this document as "
                "unsupported rather than storing near-empty content as if it had parsed."
            )
        else:
            warnings.append("OCR is disabled by configuration (ocr_enabled=false).")
        return text, ParseStatus.UNSUPPORTED_SCANNED, warnings, page_boundaries, None

    # A document can clear the average and still be part scan. Judging
    # scannedness only on the whole-document average silently loses those
    # pages: a real 18-page deck in this project's own test corpus has 13
    # pages with no text layer, but the 5 pages that do have one lift the
    # average to ~138 chars/page, so it parsed as a clean SUCCESS with
    # 72% of its content missing and no warning at all. Nothing
    # downstream could detect that -- the chunks simply did not exist,
    # and an assessor would have read "no evidence found" for content
    # that was sitting in the document.
    textless = [i for i, page in enumerate(page_texts) if len(page) < _MIN_CHARS_PER_PAGE]
    if textless and settings.ocr_enabled:
        recovered = _ocr_missing_pages(content, page_texts, textless, settings)
        if recovered is not None:
            page_texts, recovered_pages, ocr_warnings = recovered
            text = "\n\n".join(page_texts)
            # Recomputed from the spliced pages, never carried over: an
            # OCR'd page is a different length than the empty one it
            # replaced, so reusing the old boundaries would misattribute
            # every citation after the first recovered page (ADR-0042).
            page_boundaries = _page_boundaries_for(page_texts)
            warnings.extend(ocr_warnings)
            warnings.append(
                f"{len(recovered_pages)} of {page_count} pages had no text layer and were "
                f"read by local OCR (page(s) "
                f"{', '.join(str(i + 1) for i in recovered_pages[:10])}"
                f"{', and others' if len(recovered_pages) > 10 else ''}). Text from those "
                "pages is approximate -- verify any passage against the source page before "
                "relying on it as evidence."
            )
            # SUCCESS_OCR, not SUCCESS: part of this document's text came
            # from the recogniser, and a reviewer has to be told that.
            # Deliberately not a new fourth status -- the meaningful
            # distinction for a reader is "OCR was involved, treat quotes
            # as approximate", which SUCCESS_OCR already carries, and the
            # warning above says exactly which pages.
            return text, ParseStatus.SUCCESS_PARTIAL_OCR, warnings, page_boundaries, None

    return text, ParseStatus.SUCCESS, warnings, page_boundaries, None


def parse_docx(content: bytes) -> ParseResult:
    warnings: list[str] = []
    try:
        _, ceiling_warning = _zip_bomb_ceiling_warning(content, "DOCX")
    except zipfile.BadZipFile as exc:
        return "", ParseStatus.FAILED, [f"Could not open DOCX: {exc}"], None, None
    if ceiling_warning:
        return "", ParseStatus.FAILED, [ceiling_warning], None, None

    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        return "", ParseStatus.FAILED, [f"Could not open DOCX: {exc}"], None, None

    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name if para.style else "") or ""
        if style_name.lower().startswith("heading"):
            lines.append(f"# {text}")
        else:
            lines.append(text)

    full_text = "\n".join(lines)
    if not full_text.strip():
        warnings.append("DOCX contained no extractable paragraph text.")
        return full_text, ParseStatus.EMPTY, warnings, None, None

    return full_text, ParseStatus.SUCCESS, warnings, None, None


def parse_plain_text(content: bytes) -> ParseResult:
    warnings: list[str] = []
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        warnings.append(
            "Content is not valid UTF-8; decoded with latin-1 as a fallback. "
            "Review this document for encoding issues before trusting extracted evidence."
        )
        text = content.decode("latin-1", errors="replace")
        if not text.strip():
            return "", ParseStatus.ENCODING_FAILURE, warnings, None, None

    if not text.strip():
        return text, ParseStatus.EMPTY, ["File contained no text content."], None, None

    return text, ParseStatus.SUCCESS, warnings, None, None


def _render_tabular_rows(
    header: list[str], rows: list[list[str]], start_row_number: int
) -> list[tuple[str, int]]:
    """Renders each data row as "Row <N>: col1: val1 | col2: val2 | ..."
    -- self-describing (a chunk containing just "Firewall-01 | NetOps"
    with no column context would be far weaker evidence than one
    including "Asset Name: Firewall-01 | Owner: NetOps") and citable at
    the row level directly in the chunk text itself, since the existing
    chunking pipeline (chunking.py) operates on char offsets over a flat
    text string, not a separate row-index field. zip(..., strict=False)
    handles ragged rows (shorter/longer than the header) without
    crashing, a real condition in real-world spreadsheets. Blank rows
    (every cell empty) are skipped, not rendered as an empty citation.

    Returns (line_text, row_number) pairs, not bare strings (Sprint 18,
    ADR-0052) -- callers building row_boundaries need the real row_number
    per line, which is NOT the same as its position in the returned list
    once blank rows are skipped (enumerate(rows) alone would silently
    misnumber every row after the first skip).
    """
    lines: list[tuple[str, int]] = []
    for offset, row in enumerate(rows):
        if not any(cell.strip() for cell in row):
            continue
        pairs = " | ".join(f"{h}: {v}" for h, v in zip(header, row, strict=False) if h.strip())
        if pairs:
            row_number = start_row_number + offset
            lines.append((f"Row {row_number}: {pairs}", row_number))
    return lines


def _append_rendered_rows(
    rendered: list[tuple[str, int]],
    sheet_name: str | None,
    lines: list[str],
    row_boundaries: list[tuple[int, int, int, str | None]],
    cursor: int,
) -> int:
    """Appends each rendered (line_text, row_number) to `lines`, recording
    its real char offset in `row_boundaries` as it goes, and returns the
    updated cursor -- shared by parse_xlsx (per sheet) and parse_csv (one
    call, sheet_name=None), so the char-offset bookkeeping (mirroring
    parse_pdf's page_boundaries pattern above) exists in exactly one place.
    """
    for line_text, row_number in rendered:
        start = cursor
        end = start + len(line_text)
        row_boundaries.append((start, end, row_number, sheet_name))
        lines.append(line_text)
        cursor = end + 1  # the "\n" separator between lines
    return cursor


def parse_xlsx(content: bytes) -> ParseResult:
    warnings: list[str] = []
    try:
        _, ceiling_warning = _zip_bomb_ceiling_warning(content, "XLSX")
    except zipfile.BadZipFile as exc:
        return "", ParseStatus.FAILED, [f"Could not open XLSX: {exc}"], None, None
    if ceiling_warning:
        return "", ParseStatus.FAILED, [ceiling_warning], None, None

    try:
        workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        return "", ParseStatus.FAILED, [f"Could not open XLSX: {exc}"], None, None

    # Sheet names rendered as "# Sheet Name" headings so
    # chunking.py's existing structure-aware chunker (already keyed off
    # literal "# " markers, per DOCX's own heading convention above)
    # splits by sheet automatically -- no new chunking logic needed.
    lines: list[str] = []
    row_boundaries: list[tuple[int, int, int, str | None]] = []
    cursor = 0
    for sheet in workbook.worksheets:
        sheet_rows = list(sheet.iter_rows(values_only=True))
        if not sheet_rows:
            continue
        header = [str(c) if c is not None else "" for c in sheet_rows[0]]
        data_rows = [[str(c) if c is not None else "" for c in row] for row in sheet_rows[1:]]
        rendered = _render_tabular_rows(header, data_rows, start_row_number=2)
        if rendered:
            heading = f"# {sheet.title}"
            lines.append(heading)
            cursor += len(heading) + 1  # the "\n" after the heading line
            cursor = _append_rendered_rows(rendered, sheet.title, lines, row_boundaries, cursor)
    workbook.close()

    text = "\n".join(lines)
    if not text.strip():
        return (
            text,
            ParseStatus.EMPTY,
            ["XLSX contained no data rows across any sheet."],
            None,
            None,
        )

    return text, ParseStatus.SUCCESS, warnings, None, row_boundaries


def parse_csv(content: bytes) -> ParseResult:
    warnings: list[str] = []
    try:
        text_content = content.decode("utf-8")
    except UnicodeDecodeError:
        warnings.append(
            "Content is not valid UTF-8; decoded with latin-1 as a fallback. "
            "Review this document for encoding issues before trusting extracted evidence."
        )
        text_content = content.decode("latin-1", errors="replace")

    try:
        rows = list(csv.reader(io.StringIO(text_content)))
    except Exception as exc:  # the stdlib csv module can raise on pathological content
        return "", ParseStatus.FAILED, [f"Could not parse CSV: {exc}"], None, None

    if not rows:
        return "", ParseStatus.EMPTY, ["CSV contained no rows."], None, None

    rendered = _render_tabular_rows(rows[0], rows[1:], start_row_number=2)
    lines: list[str] = []
    row_boundaries: list[tuple[int, int, int, str | None]] = []
    # CSV has no sheet concept -- sheet_name is always None here.
    _append_rendered_rows(rendered, None, lines, row_boundaries, cursor=0)
    text = "\n".join(lines)
    if not text.strip():
        return text, ParseStatus.EMPTY, ["CSV contained no data rows."], None, None

    return text, ParseStatus.SUCCESS, warnings, None, row_boundaries


_PARSERS = {
    FileType.PDF: parse_pdf,
    FileType.DOCX: parse_docx,
    FileType.TXT: parse_plain_text,
    FileType.MARKDOWN: parse_plain_text,
    FileType.XLSX: parse_xlsx,
    FileType.CSV: parse_csv,
}

_EXTENSION_TO_FILE_TYPE = {
    "pdf": FileType.PDF,
    "docx": FileType.DOCX,
    "txt": FileType.TXT,
    "md": FileType.MARKDOWN,
    "xlsx": FileType.XLSX,
    "csv": FileType.CSV,
}


def file_type_from_extension(filename: str) -> FileType:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in _EXTENSION_TO_FILE_TYPE:
        raise ValueError(f"Unsupported file extension: .{ext}")
    return _EXTENSION_TO_FILE_TYPE[ext]


def parse_document(
    filename: str,
    content: bytes,
    submitter: str | None = None,
    settings: Settings | None = None,
) -> ParsedDocument:
    """Dispatch to the correct parser and wrap the result as a ParsedDocument.

    This is the single entry point services/ingestion_service.py calls. It
    never raises for a malformed document (see parser functions above) —
    it returns a status the caller must handle explicitly instead.

    `settings` is only consulted by the PDF path, for OCR (ADR-0055).
    Optional so existing callers and tests keep working unchanged; the
    real ingestion path passes the app's own Settings.
    """
    file_type = file_type_from_extension(filename)

    if not _content_matches_file_type(content, file_type):
        text, status, warnings, page_boundaries, row_boundaries = (
            "",
            ParseStatus.FAILED,
            [
                f"File content does not match its .{filename.rsplit('.', 1)[-1].lower()} extension "
                "(failed a magic-byte/binary-content check). Rejected before parsing, not passed "
                "to a format-specific parser expecting content it doesn't actually contain."
            ],
            None,
            None,
        )
    elif file_type == FileType.PDF:
        # The only parser that takes settings (OCR configuration).
        text, status, warnings, page_boundaries, row_boundaries = parse_pdf(content, settings)
    else:
        parser = _PARSERS[file_type]
        text, status, warnings, page_boundaries, row_boundaries = parser(content)

    metadata = SourceDocumentMetadata(
        document_id=_new_document_id(),
        filename=filename,
        file_type=file_type,
        submitter=submitter,
        content_hash=_content_hash(content),
        parser_version=_parser_version(file_type, status=status),
    )

    return ParsedDocument(
        metadata=metadata,
        raw_text=text,
        parse_status=status,
        parse_warnings=warnings,
        page_boundaries=page_boundaries,
        row_boundaries=row_boundaries,
    )
