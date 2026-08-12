"""Shared pytest fixtures for both src/**/tests (unit) and tests/ (integration).

Fixtures generate binary test documents (DOCX, PDF) at test time rather
than committing binary fixtures to the repo, per notebooks/README.md's
principle that anything decision-relevant should be reproducible, not
an opaque checked-in artifact.
"""

from __future__ import annotations

import io

import openpyxl
import pytest
from docx import Document as DocxDocument
from fpdf import FPDF

SAMPLE_HEADING_TEXT = "Sample Policy"
SAMPLE_BODY_TEXT = "This is a synthetic sentence used only for pipeline testing."


@pytest.fixture
def sample_docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_heading(SAMPLE_HEADING_TEXT, level=1)
    doc.add_paragraph(SAMPLE_BODY_TEXT)
    doc.add_heading("Second Section", level=1)
    doc.add_paragraph("A second synthetic paragraph for chunk-boundary testing purposes only.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, SAMPLE_BODY_TEXT * 5)
    return bytes(pdf.output())


@pytest.fixture
def scanned_like_pdf_bytes() -> bytes:
    """A syntactically valid PDF with an essentially empty page: a scan
    with nothing recoverable on it.

    Exercises the UNSUPPORTED_SCANNED detection heuristic in
    services/document_parsers.py, and, since ADR-0055 added OCR, also the
    case where OCR runs and legitimately finds nothing. For a scan with
    real readable content, use `scanned_image_pdf_bytes`.
    """
    pdf = FPDF()
    pdf.add_page()
    return bytes(pdf.output())


SCANNED_PDF_LINES = (
    "All privileged accounts are reviewed quarterly.",
    "Incidents are escalated within one hour.",
)


@pytest.fixture
def scanned_image_pdf_bytes() -> bytes:
    """A PDF whose text exists ONLY as pixels -- no text layer at all.

    Distinct from `scanned_like_pdf_bytes` (an empty page, which stands in
    for a scan with nothing recoverable): this one has real, readable
    content that OCR should be able to recover, which is what makes it a
    test of OCR rather than of the detection heuristic. Generated at test
    time for the same reason as every other fixture here -- reproducible,
    not an opaque checked-in binary.
    """
    from PIL import Image, ImageDraw, ImageFont

    image = Image.new("RGB", (1240, 400), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 34)
    except OSError:  # pragma: no cover - depends on the host's installed fonts
        font = ImageFont.load_default(size=34)
    for index, line in enumerate(SCANNED_PDF_LINES):
        draw.text((40, 40 + index * 90), line, fill="black", font=font)

    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)

    pdf = FPDF()
    pdf.add_page()
    pdf.image(buffer, x=10, y=10, w=190)
    return bytes(pdf.output())


@pytest.fixture
def invalid_utf8_bytes() -> bytes:
    return b"Some text with an invalid byte: \xff\xfe more text after it"


@pytest.fixture
def sample_xlsx_bytes() -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Assets"
    ws.append(["Asset Name", "Owner", "Criticality"])
    ws.append(["Firewall-01", "NetOps", "High"])
    ws.append(["Switch-12", "NetOps", "Medium"])
    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_csv_bytes() -> bytes:
    return (
        b"Asset Name,Owner,Criticality\n"
        b"Firewall-01,NetOps,High\n"
        b"Switch-12,NetOps,Medium\n"
    )
